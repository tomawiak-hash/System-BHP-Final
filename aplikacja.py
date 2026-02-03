import streamlit as st
import pandas as pd
import datetime
import re
import zipfile
from io import BytesIO
from docxtpl import DocxTemplate 
from docx import Document        
import google.generativeai as genai

# --- IMPORTY Z MODUŁÓW ---
from data_manager import wczytaj_liste_zawodow_lokalnie, pobierz_opis_zawodu_lokalnie
from logic_ai import generuj_kompletne_szkolenie, generuj_cel_szkolenia, generuj_test_bhp, przypisz_godziny_do_tematow, MODEL_NAME, przeprowadz_audyt_tresci
from logic_docs import generuj_dokument_z_tabela, generuj_docx_prosty, generuj_docx_z_markdown
from utils import rozplanuj_zajecia

# ----- Konfiguracja Aplikacji
st.set_page_config(page_title="Inteligentny Generator Szkoleń BHP", page_icon="🎓", layout="wide")

try:
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
except Exception:
    pass

# ----- Inicjalizacja "pamięci" aplikacji (session_state)
if 'etap' not in st.session_state:
    st.session_state.etap = 1
if 'finalna_tresc' not in st.session_state:
    st.session_state.finalna_tresc = ""
if 'zapisana_firma' not in st.session_state:
    st.session_state.zapisana_firma = ""
if 'wybrany_zawod' not in st.session_state:
    st.session_state.wybrany_zawod = ""
if 'opis_zawodu' not in st.session_state:
    st.session_state.opis_zawodu = ""
if 'spis_tresci_do_tematyki' not in st.session_state:
    st.session_state.spis_tresci_do_tematyki = []
if 'cel_szkolenia_text' not in st.session_state:
    st.session_state.cel_szkolenia_text = ""
if 'tematyka_z_godzinami' not in st.session_state:
    st.session_state.tematyka_z_godzinami = []
if 'cached_test_content' not in st.session_state:
    st.session_state.cached_test_content = None
if 'dane_do_audytu' not in st.session_state:
    st.session_state.dane_do_audytu = ""

# ----- Główny interfejs aplikacji
st.title("🎓 Inteligentny Generator Szkoleń BHP")

# =========================================================
# ETAP 1: KONFIGURACJA
# =========================================================
if st.session_state.etap == 1:
    st.header("Krok 1: Konfiguracja Szkolenia Wstępnego")
    st.info("Wprowadź dane, aby AI mogło stworzyć spersonalizowany program instruktażu stanowiskowego.")
    
    col1, col2 = st.columns(2)
    
    with col1:
        lista_zawodow = wczytaj_liste_zawodow_lokalnie()
        wybrany_zawod_nazwa = st.selectbox("Stanowisko pracy:", options=list(lista_zawodow.keys()), index=None, placeholder="Wybierz zawód...")
        nazwa_firmy = st.text_input("Nazwa firmy:", value="Przykładowa Firma S.A.")

    with col2:
        # 1. Definiujemy jedną główną listę
        LISTA_SRODOWISK = [
            "Biuro (administracja)", "Magazyn", "Praca zdalna/hybrydowa", 
            "Archiwum", "Hala produkcyjna", "Teren zewnętrzny", 
            "Wyjazdy służbowe (samochód)", "Serwerownia", "Sklep/Handel", 
            "Warsztat", "Laboratorium", "Recepcja", "Teren otwarty/Budowa"
        ]

        # 2. Wybór głównego środowiska
        srodowisko_glowne = st.selectbox(
            "Główne środowisko pracy:",
            options=LISTA_SRODOWISK,
            index=None,
            placeholder="Wybierz główne miejsce..."
        )
        
        # 3. Logika dla dodatkowych środowisk
        srodowiska_dodatkowe = [] 

        if srodowisko_glowne:
            opcje_dla_dodatkowych = [env for env in LISTA_SRODOWISK if env != srodowisko_glowne]
            srodowiska_dodatkowe = st.multiselect(
                "Dodatkowe środowisko pracy (opcjonalnie):",
                options=opcje_dla_dodatkowych,
                placeholder="Wybierz dodatkowe miejsca..."
            )
        else:
            st.multiselect(
                "Dodatkowe środowisko pracy:",
                options=[],
                disabled=True,
                placeholder="Najpierw wybierz środowisko główne ⬆️"
            )
        
    # NOWE POLE: OBOWIĄZKI
    obowiazki = st.text_area(
        "Główne obowiązki na stanowisku (Kluczowe dla Instruktażu Stanowiskowego, opcjonalne):",
        placeholder="Np. obsługa komputera, kontakt z klientem, archiwizacja dokumentów, obsługa niszczarki...",
        height=100
    )

    dodatkowe_zagrozenia = st.text_area(
        "Specyficzne zagrożenia (opcjonalnie):", 
        help="Jeśli pole zostanie puste, AI samo zidentyfikuje zagrożenia na podstawie obowiązków.",
        placeholder="Np. stres, praca przy monitorze >4h, dźwiganie pudeł z papierem..."
    )

    if st.button("🚀 Generuj kompletne szkolenie"):
        if not wybrany_zawod_nazwa:
            st.warning("Proszę wybrać zawód z listy.")
        elif not srodowisko_glowne:
            st.warning("Proszę wybrać główne środowisko pracy.")
        else:
            with st.spinner(f"Tworzenie materiałów dla: {wybrany_zawod_nazwa}..."):
                kod_zawodu = lista_zawodow[wybrany_zawod_nazwa]
                opis_zawodu = pobierz_opis_zawodu_lokalnie(kod_zawodu)
                
                # Łączenie środowisk
                srodowisko_full = srodowisko_glowne
                if srodowiska_dodatkowe:
                    lista_dodatkowych = ", ".join(srodowiska_dodatkowe)
                    srodowisko_full += f" oraz okresowo: {lista_dodatkowych}"
                
                if "Błąd:" in opis_zawodu:
                    st.error(opis_zawodu)
                else:
                    # Generowanie treści przez AI
                    finalna_tresc = generuj_kompletne_szkolenie(
                        nazwa_firmy, 
                        wybrany_zawod_nazwa, 
                        opis_zawodu, 
                        dodatkowe_zagrozenia,
                        obowiazki,
                        srodowisko_full
                    )
                
                # Zmieniamy warunek: sprawdzamy czy tekst to DOKŁADNIE komunikat błędu
                # lub czy zaczyna się od frazy błędu z logic_ai.py
                if not finalna_tresc.startswith("Błąd generowania"):
                    # Zapisujemy główne dane
                    st.session_state.finalna_tresc = finalna_tresc
                    st.session_state.zapisana_firma = nazwa_firmy
                    st.session_state.wybrany_zawod = wybrany_zawod_nazwa
                    st.session_state.dane_do_audytu = f"{obowiazki} {dodatkowe_zagrozenia}"
                    
                    # 2. Generowanie Celu Szkolenia
                    st.session_state.cel_szkolenia_text = generuj_cel_szkolenia(f"Szkolenie BHP: {wybrany_zawod_nazwa}")
                    
                    # 3. Wyciąganie spisu treści
                    st.session_state.spis_tresci_do_tematyki = re.findall(r"^(?:\d+)\.\s.*", finalna_tresc, re.MULTILINE)

                    # 4. Generowanie Tematyki z godzinami
                    st.session_state.tematyka_z_godzinami = przypisz_godziny_do_tematow(st.session_state.spis_tresci_do_tematyki)

                    # Przejście dalej
                    st.session_state.etap = 2
                    st.rerun()
                else:
                    # Tutaj trafi tylko prawdziwy błąd techniczny
                    st.error(finalna_tresc)

# =========================================================
# ETAP 2: EDYCJA I WERYFIKACJA
# =========================================================
elif st.session_state.etap == 2:
    st.header("✅ Krok 2: Weryfikacja i Edycja Treści")
    st.success("Szkolenie wygenerowane pomyślnie!")

# === AUDYT JAKOŚCI (WIDOCZNY TYLKO W TRYBIE ADMIN) ===
    # Aby zobaczyć audyt, musisz dodać do adresu strony w przeglądarce: ?tryb=admin
    # Np. localhost:8501/?tryb=admin
    
    query_params = st.query_params
    # Sprawdzamy czy w linku jest parametr "tryb" i czy ma wartość "admin"
    czy_tryb_admin = query_params.get("tryb") == "admin"

    if czy_tryb_admin:
        st.info("🔓 Tryb Administratora: Audyt Jakości jest widoczny")
        with st.expander("🔍 Raport Automatycznej Kontroli Jakości (Audyt Prawny)", expanded=True):
            st.markdown("System przeanalizował wygenerowany tekst pod kątem wymogów formalnych:")
            
            dane_input = st.session_state.get('dane_do_audytu', '')
            
            # Tu wywołujemy funkcję z logic_ai.py (musisz mieć ją zaimportowaną)
            wyniki = przeprowadz_audyt_tresci(st.session_state.finalna_tresc, dane_input)
            
            for kategoria, status in wyniki.items():
                c1, c2 = st.columns([0.7, 0.3])
                c1.write(f"**{kategoria}**")
                
                if status == "SKIP":
                    c2.caption("⚪ Brak danych (Pominięto)")
                elif status is True:
                    c2.success("✅ OK")
                else:
                    c2.error("❌ BRAK")

    st.markdown("---")

    # 1. EDYTOR HARMONOGRAMU
    st.subheader("🛠️ Harmonogram Szkolenia (Edycja Godzin)")
    st.info("Program ramowy jest stały. Możesz dostosować jedynie liczbę godzin dla poszczególnych bloków.")

    if st.session_state.tematyka_z_godzinami:
        df = pd.DataFrame(st.session_state.tematyka_z_godzinami)
        
        column_config = {
            "nazwa": st.column_config.TextColumn(
                "Temat (Zgodny z Ramowym Programem)", 
                width="large", 
                disabled=True,
                help="Nazwy tematów wynikają z rozporządzenia i nie mogą być zmieniane."
            ),
            "godziny": st.column_config.NumberColumn(
                "Godziny (45min)", 
                min_value=0.1, 
                max_value=4.0,
                step=0.1, 
                format="%.1f h",
                help="Wpisz wartość od 0.1 do 4.0 h"
            )
        }

        edited_df = st.data_editor(
            df, 
            column_config=column_config, 
            use_container_width=True,
            num_rows="fixed", 
            key="editor_tematyki",
            hide_index=True
        )
        
        st.session_state.tematyka_z_godzinami = edited_df.to_dict('records')
        total_h = edited_df['godziny'].sum()
        
        col_sum, col_warn = st.columns([1, 3])
        col_sum.caption(f"📊 Razem: **{total_h:.1f} h**")
        
        if total_h < 5.0:
            col_warn.warning("⚠️ Suma godzin jest niska. Upewnij się, że spełniasz wymogi min. 3h ogólnego + min. 2h stanowiskowego.")
        elif total_h > 16.0:
            col_warn.warning("⚠️ Bardzo duża liczba godzin (ponad 2 dni szkolenia). Sprawdź poprawność.")

    st.markdown("---")

    # 2. PODGLĄD TREŚCI
    st.subheader("📖 Treść Szkolenia")
    
    with st.expander("✏️ Kliknij tutaj, aby ręcznie edytować tekst źródłowy"):
        st.text_area("Edycja treści:", value=st.session_state.finalna_tresc, height=300, key="edycja_tekstu_area")
        if st.session_state.edycja_tekstu_area != st.session_state.finalna_tresc:
            st.session_state.finalna_tresc = st.session_state.edycja_tekstu_area
            st.rerun()

    with st.expander("📄 Podgląd sformatowanej treści szkolenia (Kliknij, aby zwinąć/rozwinąć)", expanded=True):
        st.markdown(st.session_state.finalna_tresc, unsafe_allow_html=True)

    st.markdown("---")


    # 3. PRZYCISKI NAWIGACJI
    col_btn1, col_btn2 = st.columns([1, 1])
    
    with col_btn1:
        docx_file = generuj_docx_z_markdown(st.session_state.finalna_tresc)
        st.download_button(
            label="📥 Pobierz treść jako WORD (.docx)",
            data=docx_file,
            file_name=f"Szkolenie_{st.session_state.wybrany_zawod}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with col_btn2:
        if st.button("📄 Zatwierdź i przejdź do dokumentów", type="primary", use_container_width=True):
            st.session_state.etap = 3
            st.rerun()
            
    if st.button("🔙 Wróć do wyboru zawodu", type="secondary"):
        st.session_state.etap = 1
        st.rerun()

# =========================================================
# ETAP 3: GENERATOR DOKUMENTACJI
# =========================================================
elif st.session_state.etap == 3:
    st.header("✅ Krok 3: Generator Dokumentacji")
    
    # --- SEKCJA DANYCH WSPÓLNYCH ---
    with st.container(border=True):
        st.subheader("🛠️ Konfiguracja danych")
        
        # 1. UCZESTNICY
        st.markdown("**Lista uczestników** \n*Wpisz tylko: Imię Nazwisko, Data Urodzenia*", unsafe_allow_html=True)
        uczestnicy_input = st.text_area(
            label="Lista uczestników", 
            label_visibility="collapsed",
            height=100, 
            key="uczestnicy_lista_input", 
            placeholder="Jan Kowalski, 12.05.1985\nAnna Nowak, 20.01.1990"
        )
        
        uczestnicy_dane_lista = []
        bledne_linie_detale = []
        
        if uczestnicy_input:
            lines = uczestnicy_input.strip().splitlines()
            for i, linia in enumerate(lines):
                linia_clean = linia.strip()
                if not linia_clean: continue 
                
                czesci = [c.strip() for c in linia_clean.split(',')]
                
                if len(czesci) != 2:
                    bledne_linie_detale.append(f"❌ Linia {i+1}: Nieprawidłowy format. Wymagane: 'Imię Nazwisko, Data'.")
                    continue

                data_raw = czesci[1]
                if not re.match(r"^\d{2}\.\d{2}\.\d{4}$", data_raw):
                    bledne_linie_detale.append(f"❌ Linia {i+1}: Zły format daty '{data_raw}'. Wymagane DD.MM.RRRR.")
                    continue
                
                uczestnicy_dane_lista.append({
                    'index': len(uczestnicy_dane_lista) + 1, 
                    'imie_nazwisko': czesci[0], 
                    'miejsce_pracy': st.session_state.zapisana_firma,
                    'funkcja': st.session_state.wybrany_zawod,
                    'data_urodzenia': czesci[1], 
                    'ocena': '', 
                    'uwagi': ''
                })

        if bledne_linie_detale:
            st.error(f"Znaleziono błędy w {len(bledne_linie_detale)} wierszach:")
            for blad in bledne_linie_detale:
                st.text(blad)
        
        if uczestnicy_dane_lista:
            with st.expander(f"✅ Poprawnie wczytano {len(uczestnicy_dane_lista)} uczestników", expanded=False):
                st.dataframe(pd.DataFrame(uczestnicy_dane_lista)[['imie_nazwisko', 'miejsce_pracy', 'funkcja', 'data_urodzenia']], use_container_width=True, hide_index=True)

        st.markdown("---")

        # 2. BAZA KADRY
        st.markdown("### ⚙️ Baza Kadry")
        col_kadra1, col_kadra2 = st.columns(2)

        with col_kadra1:
            if 'baza_wykladowcow_text' not in st.session_state:
                st.session_state.baza_wykladowcow_text = "Jan Nowak, Firma BHP, Specjalista BHP\nAnna Kowalska, Firma Med, Ratownik"
            st.markdown("**1. Baza Wykładowców**", unsafe_allow_html=True)
            baza_wykladowcow = st.text_area("Baza Wykładowców", label_visibility="collapsed", value=st.session_state.baza_wykladowcow_text, height=150, key="baza_wykladowcow_key")
            opcje_wykladowcow = [x.strip() for x in baza_wykladowcow.splitlines() if x.strip()]

        with col_kadra2:
            if 'baza_komisji_text' not in st.session_state:
                st.session_state.baza_komisji_text = "Jan Nowak, Firma BHP, Przewodniczący\nAnna Kowalska, Firma BHP, Członek Komisji"
            st.markdown("**2. Baza Komisji Egzaminacyjnej**", unsafe_allow_html=True)
            baza_komisji = st.text_area("Baza Komisji", label_visibility="collapsed", value=st.session_state.baza_komisji_text, height=150, key="baza_komisji_key")
            opcje_komisji = [x.strip() for x in baza_komisji.splitlines() if x.strip()]

        st.markdown("---")

        # 3. HARMONOGRAM
        st.markdown("### 🗓️ Harmonogram Szkolenia")
        col_d1, col_d2 = st.columns(2)
        dzisiaj = datetime.date.today()
        
        with col_d1:
            data_start = st.date_input("Data rozpoczęcia:", key="doc_data_start", value=dzisiaj)
            nr_kursu = st.text_input("Numer kursu:", "01/BHP/2025", key="doc_nr_kursu")
            kierownik_kursu = st.text_input("Kierownik kursu:", "Anna Kowalska", key="doc_kierownik")
        
        sugerowana_data_koniec = data_start
        if st.session_state.tematyka_z_godzinami:
            _, obliczona_data = rozplanuj_zajecia(st.session_state.tematyka_z_godzinami, data_start)
            if obliczona_data > data_start: sugerowana_data_koniec = obliczona_data

        with col_d2:
            wartosc_domyslna = sugerowana_data_koniec
            if 'doc_data_koniec' in st.session_state:
                if st.session_state.doc_data_koniec >= data_start: wartosc_domyslna = st.session_state.doc_data_koniec
                else: wartosc_domyslna = data_start

            data_koniec = st.date_input("Data zakończenia:", key="doc_data_koniec", value=wartosc_domyslna, min_value=data_start)
            miejscowosc = st.text_input("Miejscowość:", "Łódź", key="doc_miejscowosc")
            
            wartosc_domyslna_wyst = data_koniec
            if 'doc_data_wyst' in st.session_state and st.session_state.doc_data_wyst >= data_koniec:
                 wartosc_domyslna_wyst = st.session_state.doc_data_wyst

            data_wystawienia = st.date_input("Data wystawienia dokumentów:", key="doc_data_wyst", value=wartosc_domyslna_wyst, min_value=data_koniec)

        st.session_state.faktyczna_data_koniec = data_koniec

    st.write("") 

    tab1, tab2, tab3 = st.tabs(["📜 Zaświadczenia i Rejestr", "📅 Przebieg Szkolenia", "📝 Wykaz i Pytania"])

    # --- TAB 1 ---
    with tab1:
        st.info("Dokumentacja potwierdzająca odbycie instruktażu ogólnego i stanowiskowego.")
        col_z1, col_z2 = st.columns(2)
        
        with col_z1:
            st.subheader("📄 Karta Szkolenia Wstępnego")
            with st.container(border=True):
                instruktor_ogolny = st.selectbox("Instruktor (Instruktaż Ogólny):", options=opcje_wykladowcow, index=0 if opcje_wykladowcow else None, key="inst_ogolny_sel")
                instruktor_stanowiskowy = st.selectbox("Instruktor (Instruktaż Stanowiskowy):", options=opcje_wykladowcow, index=0 if opcje_wykladowcow else None, key="inst_stan_sel")
                
                data_stanowiskowego = st.date_input(
                    "Data instruktażu stanowiskowego:", 
                    value=data_koniec, 
                    min_value=data_start,
                    key="date_stanowiskowy_input"
                )

                st.markdown("---")
                wybrany_uczestnik = st.selectbox("Wybierz uczestnika do podglądu:", options=[u['imie_nazwisko'] for u in uczestnicy_dane_lista], index=None, key="sel_uczestnik_karta")
                
                if st.button("Generuj Kartę Szkolenia (Pojedynczą)", use_container_width=True, key="btn_gen_karta_single"):
                    if wybrany_uczestnik and instruktor_ogolny and instruktor_stanowiskowy:
                        osoba = next((u for u in uczestnicy_dane_lista if u['imie_nazwisko'] == wybrany_uczestnik), None)
                        inst_ogolny_nazwisko = instruktor_ogolny.split(',')[0].strip()
                        inst_stan_nazwisko = instruktor_stanowiskowy.split(',')[0].strip()

                        context = {
                            'nazwa_firmy': st.session_state.zapisana_firma,
                            'imie_nazwisko': osoba['imie_nazwisko'],
                            'komorka_organizacyjna': osoba['miejsce_pracy'], 
                            'stanowisko': osoba['funkcja'],
                            'dzien_rozpoczecia': data_start.strftime("%d.%m.%Y"),
                            'instruktor_ogolny': inst_ogolny_nazwisko,
                            'data_stanowiskowego': data_stanowiskowego.strftime("%d.%m.%Y"),
                            'instruktor_stanowiskowy': inst_stan_nazwisko
                        }
                        plik = generuj_docx_prosty("Wzor-Karta-szkolenia-wstepnego-BHP.docx", context, "Karta.docx")
                        if plik:
                            st.download_button("📥 Pobierz Kartę", plik, f"Karta_Szkolenia_{osoba['imie_nazwisko']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key=f"dl_karta_{osoba['index']}")
                    else:
                        st.warning("Uzupełnij instruktorów i wybierz uczestnika.")

        with col_z2:
            st.subheader("📚 Rejestr Zaświadczeń")
            with st.container(border=True):
                st.write(f"Rejestr dla kursu: **{nr_kursu}**")
                
                if st.button("Generuj Rejestr", use_container_width=True, key="btn_gen_rejestr_final"):
                    if uczestnicy_dane_lista:
                        rejestr_dane = []
                        for i, u in enumerate(uczestnicy_dane_lista):
                            rejestr_dane.append({'numer': f"{nr_kursu}/{i+1}", 'imie_nazwisko': u['imie_nazwisko'], 'uwagi': ''})
                        
                        context = {
                            'rodzaj_szkolenia': "wstępnego", 'nr_kursu': nr_kursu,
                            'kierownik_nazwisko': kierownik_kursu,
                            'data_wystawienia': data_wystawienia.strftime("%d.%m.%Y"),
                            'nazwa_organizatora': st.session_state.zapisana_firma, 'miejsce': miejscowosc
                        }
                        plik, blad = generuj_dokument_z_tabela("rejestr_zaswiadczen_szablon_uproszczony.docx", context, rejestr_dane, ['numer', 'imie_nazwisko', 'podpis_dummy', 'uwagi'], index_tabeli=2)
                        if plik: st.download_button("📥 Pobierz Rejestr", plik, "Rejestr.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="dl_rejestr_final")
                        else: st.error(blad)
                    else: st.error("Brak uczestników.")

    # --- TAB 2 ---
    with tab2:
        st.info("Dokumentacja dotycząca programu, harmonogramu i realizacji zajęć.")
        col_p1, col_p2 = st.columns(2)
        
        with col_p1:
            st.subheader("📋 Tematyka Szkolenia")
            with st.container(border=True):
                if st.button("Generuj Tematykę", use_container_width=True):
                    tematyka = st.session_state.tematyka_z_godzinami
                    if tematyka:
                        total_h = sum(float(t.get('godziny', 0)) for t in tematyka)
                        tematyka_display = [{"nazwa": t.get('nazwa',''), "godziny": t.get('godziny',0), "praktyka": "0"} for t in tematyka]
                        tematyka_display.append({"nazwa": "RAZEM:", "godziny": f"{total_h:.1f}", "praktyka": "0"})

                        plik, blad = generuj_dokument_z_tabela("tematyka_szablon_uproszczony.docx", {}, tematyka_display, ['nazwa', 'godziny', 'praktyka'])
                        if plik:
                            st.download_button("📥 Pobierz Tematykę", plik, "Tematyka.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                        else:
                            st.error(blad)
                    else:
                        st.error("Brak danych tematyki.")

        with col_p2:
            st.subheader("🗓️ Dziennik Zajęć")
            with st.container(border=True):
                if st.button("Generuj Dziennik Zajęć", use_container_width=True):
                    tematyka = st.session_state.tematyka_z_godzinami
                    if tematyka:
                        zajecia, faktyczna_data = rozplanuj_zajecia(tematyka, data_start)
                        st.session_state.faktyczna_data_koniec = faktyczna_data
                        
                        doc_tpl = DocxTemplate("dziennik_zajec_szablon_uproszczony.docx")
                        doc_tpl.render({'nazwa_organizatora': st.session_state.zapisana_firma})
                        bio = BytesIO(); doc_tpl.save(bio); bio.seek(0)
                        
                        doc = Document(bio)
                        if doc.tables:
                            table = doc.tables[0] 
                            for i, z in enumerate(zajecia):
                                row = table.add_row().cells
                                if len(row) >= 6:
                                    row[0].text = str(i + 1)
                                    row[1].text = z['data']
                                    row[2].text = str(z['godziny'])
                                    row[3].text = z['przedmiot']
                                    row[4].text = z['temat']
                                    row[5].text = "" 
                        
                        final_bio = BytesIO(); doc.save(final_bio); final_bio.seek(0)
                        st.download_button("📥 Pobierz Dziennik Zajęć", final_bio, "Dziennik_Zajec.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    else:
                        st.error("Brak tematyki.")

        st.subheader("📓 Dziennik Lekcyjny")
        with st.container(border=True):
            st.write("Wybierz wykładowców i przypisz im godziny.")
            wybrani_wykladowcy = st.multiselect("Wybierz wykładowców:", options=opcje_wykladowcow, key="wykladowcy_multiselect")
            lista_do_przetworzenia = wybrani_wykladowcy if wybrani_wykladowcy else opcje_wykladowcow
            
            # --- NAPRAWA BŁĘDU FLOAT (inicjalizacja 0.0) ---
            total_h = 0.0
            if st.session_state.tematyka_z_godzinami:
                total_h = sum(float(t.get('godziny', 0)) for t in st.session_state.tematyka_z_godzinami)
            
            godziny_na_osobe = {}
            if lista_do_przetworzenia:
                st.markdown("#### ⏱️ Przydział godzin")
                domyslna_wartosc = total_h / len(lista_do_przetworzenia) if len(lista_do_przetworzenia) > 0 else 0.0
                cols = st.columns(3)
                
                for i, osoba_raw in enumerate(lista_do_przetworzenia):
                    nazwisko = osoba_raw.split(',')[0].strip()
                    with cols[i % 3]:
                        # --- NAPRAWA: WYMUSZENIE FLOAT W MAX_VALUE ---
                        godziny_na_osobe[osoba_raw] = st.number_input(
                            f"{nazwisko} (h):",
                            min_value=0.0,
                            max_value=float(total_h * 2.0), # <--- TU BYŁ BŁĄD (wymuszamy float)
                            value=float(f"{domyslna_wartosc:.1f}"),
                            step=0.5,
                            key=f"godziny_wyk_{i}"
                        )
                
                suma_wpisana = sum(godziny_na_osobe.values())
                if abs(suma_wpisana - total_h) > 0.1:
                    st.warning(f"⚠️ Uwaga: Suma godzin wykładowców ({suma_wpisana:.1f}h) różni się od sumy godzin szkolenia ({total_h:.1f}h).")
                else:
                    st.success(f"✅ Suma godzin się zgadza ({total_h:.1f}h).")

            if st.button("Generuj Dziennik Lekcyjny", use_container_width=True):
                wykladowcy_lista = [] 
                if not lista_do_przetworzenia:
                    st.error("Brak wykładowców! Wpisz ich w Bazie Kadry.")
                else:
                    for linia in lista_do_przetworzenia:
                        parts = [p.strip() for p in linia.split(',', 2)]
                        if len(parts) == 3:
                            h_user = godziny_na_osobe.get(linia, 0.0)
                            wykladowcy_lista.append({
                                'imie_nazwisko': parts[0], 
                                'miejsce_pracy': parts[1], 
                                'funkcja': parts[2],
                                'przedmiot': 'Szkolenie wstępne BHP', 
                                'godziny_plan': f"{h_user:.1f}", 
                                'godziny_wykonanie': f"{h_user:.1f}"
                            })
                    
                    if wykladowcy_lista:
                        suma_h = sum(float(w['godziny_plan']) for w in wykladowcy_lista)
                        wykladowcy_lista.append({
                            'imie_nazwisko': '', 'miejsce_pracy': '', 'funkcja': '', 
                            'przedmiot': 'RAZEM:', 
                            'godziny_plan': f"{suma_h:.1f}", 
                            'godziny_wykonanie': f"{suma_h:.1f}"
                        })

                        context = {
                            'nazwa_organizatora': st.session_state.zapisana_firma,
                            'dla_kogo': f"Szkolenie dla: {st.session_state.wybrany_zawod}",
                            'data_od': data_start.strftime("%d.%m.%Y"), 
                            'data_do': data_koniec.strftime("%d.%m.%Y"),
                            'miejsce': miejscowosc, 
                            'kierownik_nazwisko': kierownik_kursu,
                            'kierownik_miejsce_pracy_funkcja': "Kierownik Szkolenia" 
                        }

                        plik, blad = generuj_dokument_z_tabela("dziennik_lekcyjny_szablon_uproszczony.docx", context, wykladowcy_lista, ['imie_nazwisko', 'miejsce_pracy', 'funkcja', 'przedmiot', 'godziny_plan', 'godziny_wykonanie'], index_tabeli=4)
                        if plik:
                            st.download_button("📥 Pobierz Dziennik Lekcyjny", plik, "Dziennik_Lekcyjny.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                        else: st.error(blad)

    # --- TAB 3 ---
    with tab3:
        col_e1, col_e2 = st.columns(2)
        with col_e1:
            st.subheader("👥 Wykaz Uczestników")
            with st.container(border=True):
                if st.button("Generuj Wykaz", use_container_width=True, key="btn_gen_wykaz_final"):
                    if uczestnicy_dane_lista:
                        plik, blad = generuj_dokument_z_tabela("wykaz_uczestnikow_szablon_uproszczony.docx", {}, uczestnicy_dane_lista, ['imie_nazwisko', 'miejsce_pracy', 'funkcja', 'data_urodzenia'])
                        if plik:
                            st.download_button("📥 Pobierz Wykaz", plik, "Wykaz_Uczestnikow.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="dl_wykaz_final")
                        else: st.error(blad)
                    else: st.warning("Brak uczestników.")

        with col_e2:
            st.subheader("❓ Pytania Kontrolne")
            with st.container(border=True):
                if st.button("Generuj Pytania Kontrolne", use_container_width=True, key="btn_gen_pytania_final"):
                    if st.session_state.finalna_tresc:
                        with st.spinner("AI opracowuje pytania sprawdzające..."):
                            tresc_pytan, _ = generuj_test_bhp(st.session_state.finalna_tresc)
                            st.session_state.cached_test_content = tresc_pytan
                    else:
                        st.warning("Najpierw wygeneruj program szkolenia w Kroku 1.")

                if st.session_state.cached_test_content:
                    st.success("Pytania gotowe.")
                    ctx_pytania = {
                             'nazwa_szkolenia': f"Szkolenie wstępne i stanowiskowe dla {st.session_state.wybrany_zawod}", 
                             'tresc_testu': st.session_state.cached_test_content
                         }
                    plik_pytania = generuj_docx_prosty("test_szablon.docx", ctx_pytania, "Pytania.docx")
                    if plik_pytania:
                        st.download_button("📥 Pobierz Arkusz Pytań", plik_pytania, "Pytania_Kontrolne.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="dl_pytania_final")

    # --- ZIP ---
    st.markdown("---")
    st.subheader("📦 Pobierz wszystko")
    st.info("Wygeneruj komplet dokumentacji jednym kliknięciem.")

    if st.button("Generuj paczkę ZIP ze wszystkimi dokumentami", type="primary", use_container_width=True, key="btn_zip_final"):
        if not uczestnicy_dane_lista:
            st.error("Brakuje listy uczestników!")
        elif not st.session_state.tematyka_z_godzinami:
            st.error("Brakuje tematyki szkolenia!")
        else:
            zip_buffer = BytesIO()
            try:
                with zipfile.ZipFile(zip_buffer, "w") as zf:
                    # Logika ZIP bez zmian - kopiujemy istniejące funkcjonalności
                    # ... (skrócone dla czytelności, logika identyczna jak w poprzedniej wersji)
                    
                    # 1. KARTY
                    inst_ogolny_zip = st.session_state.get("inst_ogolny_sel", "Instruktor")
                    inst_stan_zip = st.session_state.get("inst_stan_sel", "Kierownik")
                    data_stan_input = st.session_state.get("date_stanowiskowy_input", data_koniec)
                    i_ogolny = str(inst_ogolny_zip).split(',')[0].strip()
                    i_stan = str(inst_stan_zip).split(',')[0].strip()
                    d_stan = data_stan_input.strftime("%d.%m.%Y")

                    for u in uczestnicy_dane_lista:
                        context_karta = {'nazwa_firmy': st.session_state.zapisana_firma, 'imie_nazwisko': u['imie_nazwisko'], 'komorka_organizacyjna': u['miejsce_pracy'], 'stanowisko': u['funkcja'], 'dzien_rozpoczecia': data_start.strftime("%d.%m.%Y"), 'instruktor_ogolny': i_ogolny, 'data_stanowiskowego': d_stan, 'instruktor_stanowiskowy': i_stan}
                        plik = generuj_docx_prosty("Wzor-Karta-szkolenia-wstepnego-BHP.docx", context_karta, "temp.docx")
                        if plik: zf.writestr(f"Karty_Szkolenia/Karta_{u['imie_nazwisko']}.docx", plik.getvalue())

                    # 2. REJESTR
                    rejestr_dane = []
                    for i, u in enumerate(uczestnicy_dane_lista):
                        rejestr_dane.append({'numer': f"{nr_kursu}/{i+1}", 'imie_nazwisko': u['imie_nazwisko'], 'uwagi': ''})
                    context_rej = {'rodzaj_szkolenia': "wstępnego", 'nr_kursu': nr_kursu, 'kierownik_nazwisko': kierownik_kursu, 'data_wystawienia': data_wystawienia.strftime("%d.%m.%Y"), 'nazwa_organizatora': st.session_state.zapisana_firma, 'miejsce': miejscowosc}
                    plik, _ = generuj_dokument_z_tabela("rejestr_zaswiadczen_szablon_uproszczony.docx", context_rej, rejestr_dane, ['numer', 'imie_nazwisko', 'podpis_dummy', 'uwagi'], index_tabeli=2)
                    if plik: zf.writestr("Rejestr_Zaswiadczen.docx", plik.getvalue())

                    # 3. TEMATYKA
                    tematyka = st.session_state.tematyka_z_godzinami
                    total_h = sum(t.get('godziny', 0) for t in tematyka if isinstance(t.get('godziny'), (int, float)))
                    tematyka_display = [{"nazwa": t.get('nazwa',''), "godziny": t.get('godziny',0), "praktyka": "0"} for t in tematyka]
                    tematyka_display.append({"nazwa": "RAZEM:", "godziny": f"{total_h:.1f}", "praktyka": "0"})
                    plik, _ = generuj_dokument_z_tabela("tematyka_szablon_uproszczony.docx", {}, tematyka_display, ['nazwa', 'godziny', 'praktyka'])
                    if plik: zf.writestr("Tematyka_Szkolenia.docx", plik.getvalue())

                    # 4. DZIENNIK ZAJĘĆ
                    zajecia, _ = rozplanuj_zajecia(tematyka, data_start)
                    doc_tpl = DocxTemplate("dziennik_zajec_szablon_uproszczony.docx")
                    doc_tpl.render({'nazwa_organizatora': st.session_state.zapisana_firma})
                    bio = BytesIO(); doc_tpl.save(bio); bio.seek(0)
                    doc = Document(bio)
                    if doc.tables:
                        table = doc.tables[0]
                        for i, z in enumerate(zajecia):
                            row = table.add_row().cells
                            if len(row) >= 6:
                                row[0].text = str(i+1); row[1].text = z['data']; row[2].text = str(z['godziny']); row[3].text = z['przedmiot']; row[4].text = z['temat']
                    final_bio = BytesIO(); doc.save(final_bio); final_bio.seek(0)
                    zf.writestr("Dziennik_Zajec.docx", final_bio.getvalue())

                    # 5. DZIENNIK LEKCYJNY
                    wybrani_wykladowcy_zip = st.session_state.get("wykladowcy_multiselect", [])
                    baza_wyk_raw = st.session_state.get("baza_wykladowcow_key", "")
                    opcje_bazy_wyk = [x.strip() for x in baza_wyk_raw.splitlines() if x.strip()]
                    finalna_lista_zip = wybrani_wykladowcy_zip if wybrani_wykladowcy_zip else opcje_bazy_wyk
                    
                    if finalna_lista_zip:
                        total_h_szkolenia = sum(float(t.get('godziny', 0)) for t in tematyka)
                        liczba_wykladowcow = len(finalna_lista_zip)
                        godziny_na_glowe = total_h_szkolenia / liczba_wykladowcow if liczba_wykladowcow > 0 else 0
                        wykladowcy_lista = []
                        for linia in finalna_lista_zip:
                            parts = [p.strip() for p in linia.split(',', 2)]
                            if len(parts) == 3:
                                wykladowcy_lista.append({'imie_nazwisko': parts[0], 'miejsce_pracy': parts[1], 'funkcja': parts[2], 'przedmiot': 'Szkolenie wstępne BHP', 'godziny_plan': f"{godziny_na_glowe:.1f}", 'godziny_wykonanie': f"{godziny_na_glowe:.1f}"})
                        if wykladowcy_lista:
                            wykladowcy_lista.append({'imie_nazwisko': '', 'miejsce_pracy': '', 'funkcja': '', 'przedmiot': 'RAZEM:', 'godziny_plan': f"{total_h_szkolenia:.1f}", 'godziny_wykonanie': f"{total_h_szkolenia:.1f}"})
                            context_lek = {'nazwa_organizatora': st.session_state.zapisana_firma, 'dla_kogo': f"Szkolenie dla: {st.session_state.wybrany_zawod}", 'data_od': data_start.strftime("%d.%m.%Y"), 'data_do': data_koniec.strftime("%d.%m.%Y"), 'miejsce': miejscowosc, 'kierownik_nazwisko': kierownik_kursu, 'kierownik_miejsce_pracy_funkcja': "Kierownik Szkolenia"}
                            plik, _ = generuj_dokument_z_tabela("dziennik_lekcyjny_szablon_uproszczony.docx", context_lek, wykladowcy_lista, ['imie_nazwisko', 'miejsce_pracy', 'funkcja', 'przedmiot', 'godziny_plan', 'godziny_wykonanie'], index_tabeli=4)
                            if plik: zf.writestr("Dziennik_Lekcyjny.docx", plik.getvalue())

                    # 6. WYKAZ
                    plik, _ = generuj_dokument_z_tabela("wykaz_uczestnikow_szablon_uproszczony.docx", {}, uczestnicy_dane_lista, ['imie_nazwisko', 'miejsce_pracy', 'funkcja', 'data_urodzenia'])
                    if plik: zf.writestr("Wykaz_Uczestnikow.docx", plik.getvalue())

                    # 7. PYTANIA
                    if not st.session_state.cached_test_content and st.session_state.finalna_tresc:
                        try:
                            tresc_pytan, _ = generuj_test_bhp(st.session_state.finalna_tresc)
                            st.session_state.cached_test_content = tresc_pytan
                        except: pass
                    if st.session_state.cached_test_content:
                         ctx_pytania = {
                             'nazwa_szkolenia': f"Szkolenie wstępne i stanowiskowe dla {st.session_state.wybrany_zawod}", 
                             'tresc_testu': st.session_state.cached_test_content
                         }
                         plik = generuj_docx_prosty("test_szablon.docx", ctx_pytania, "temp.docx")
                         if plik: zf.writestr("Pytania_Kontrolne.docx", plik.getvalue())

                    # 8. TREŚĆ (Jako DOCX)
                    # Używamy funkcji konwertującej Markdown na Word (tej samej co przy przycisku pobierania)
                    docx_tresc = generuj_docx_z_markdown(st.session_state.finalna_tresc)
                    
                    # Zapisujemy do ZIP jako plik .docx (pobieramy bajty za pomocą .getvalue())
                    zf.writestr(f"Program_Szkolenia_{st.session_state.wybrany_zawod}.docx", docx_tresc.getvalue())

                zip_buffer.seek(0)
                st.success("Paczka dokumentów gotowa!")
                st.download_button(label="📦 POBIERZ PLIK ZIP", data=zip_buffer, file_name=f"Komplet_BHP_{st.session_state.wybrany_zawod}.zip", mime="application/zip", use_container_width=True, key="dl_zip_final")

            except Exception as e:
                st.error(f"Wystąpił błąd podczas tworzenia archiwum ZIP: {e}")

    st.markdown("---")
    if st.button("🔄 Zacznij od nowa (Nowe Szkolenie)", type="secondary"):
        st.session_state.etap = 1
        st.rerun()