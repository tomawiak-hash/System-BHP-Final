from docxtpl import DocxTemplate
from docx import Document
from io import BytesIO
import streamlit as st

def generuj_dokument_z_tabela(nazwa_szablonu, context, dane_tabeli=None, mapowanie_kolumn=None, index_tabeli=0):
    """Generuje dokument Word z dynamiczną tabelą."""
    try:
        doc_tpl = DocxTemplate(nazwa_szablonu)
        doc_tpl.render(context)
        
        temp_bio = BytesIO()
        doc_tpl.save(temp_bio)
        temp_bio.seek(0)

        if dane_tabeli and mapowanie_kolumn:
            doc = Document(temp_bio)
            if doc.tables and len(doc.tables) > index_tabeli:
                table = doc.tables[index_tabeli]
                for i, wiersz_dane in enumerate(dane_tabeli):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i + 1) # Kolumna 0 to zawsze Lp.
                    
                    for col_idx, klucz in enumerate(mapowanie_kolumn):
                        target_idx = col_idx + 1
                        if target_idx < len(row_cells):
                            wartosc = str(wiersz_dane.get(klucz, ''))
                            row_cells[target_idx].text = wartosc
            else:
                return None, f"Brak tabeli o indeksie {index_tabeli} w szablonie."

            final_bio = BytesIO()
            doc.save(final_bio)
            final_bio.seek(0)
            return final_bio, None
        
        return temp_bio, None

    except Exception as e:
        return None, str(e)

def generuj_docx_prosty(nazwa_szablonu, kontekst, nazwa_pliku_wynikowego):
    """Wrapper dla prostych dokumentów bez dynamicznych tabel."""
    try:
        doc = DocxTemplate(nazwa_szablonu)
        doc.render(kontekst)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
    except Exception as e:
        st.error(f"Błąd generowania pliku '{nazwa_pliku_wynikowego}': {e}")
        return None
    
import re
from docx.shared import Pt, RGBColor

def generuj_docx_z_markdown(tekst_markdown):
    """
    Konwertuje tekst Markdown (nagłówki #, pogrubienia **) na sformatowany dokument Word.
    """
    doc = Document()
    
    # Ustawiamy domyślną czcionkę
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Dzielimy tekst na akapity
    akapity = tekst_markdown.split('\n')
    
    for akapit in akapity:
        akapit = akapit.strip()
        if not akapit:
            continue
            
        # Obsługa nagłówków
        if akapit.startswith('# '):
            doc.add_heading(akapit[2:], level=1)
        elif akapit.startswith('## '):
            doc.add_heading(akapit[3:], level=2)
        elif akapit.startswith('### '):
            doc.add_heading(akapit[4:], level=3)
        else:
            # Obsługa zwykłego tekstu z pogrubieniami
            p = doc.add_paragraph()
            
            # Rozdzielamy tekst wg znaczników ** (pogrubienie)
            # Przykład: "To jest **ważne** zdanie." -> ["To jest ", "ważne", " zdanie."]
            czesci = re.split(r'(\*\*.*?\*\*)', akapit)
            
            for czesc in czesci:
                if czesc.startswith('**') and czesc.endswith('**'):
                    # To jest tekst pogrubiony
                    run = p.add_run(czesc[2:-2]) # Usuwamy gwiazdki
                    run.bold = True
                else:
                    # To jest zwykły tekst
                    p.add_run(czesc)
                    
    # Zapis do bufora
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio