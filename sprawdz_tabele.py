import streamlit as st
from docx import Document
import os

st.title("🕵️ Diagnostyka Tabel w Wordzie")

plik = "dziennik_lekcyjny_szablon_uproszczony.docx"

if st.button("Sprawdź strukturę pliku"):
    if os.path.exists(plik):
        doc = Document(plik)
        st.write(f"📄 Plik: **{plik}**")
        st.write(f"🔢 Znaleziono tabel: **{len(doc.tables)}**")
        
        for i, table in enumerate(doc.tables):
            st.markdown(f"---")
            st.subheader(f"Tabela Index: {i}")
            st.write(f"Liczba wierszy: {len(table.rows)}")
            st.write(f"Liczba kolumn: {len(table.columns)}")
            
            # Pokazujemy treść pierwszego wiersza (nagłówka)
            if len(table.rows) > 0:
                row_data = [cell.text.strip() for cell in table.rows[0].cells]
                st.code(f"Zawartość nagłówka: {row_data}")
            else:
                st.warning("Tabela jest pusta (0 wierszy).")
    else:
        st.error(f"Nie znaleziono pliku: {plik}")